"""generate_report.py - Build BaoCao_NCKH_MediSignAI_NguyenDuyThuan_2311555799.docx
Tuan thu Quy chuan HSU 4.1.1 - Cau truc 1 - Times New Roman 12 - line 1.5
Ngon ngu: Tieng Viet - APA 6 - khong dung 'toi'
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Twips

ROOT = Path(r"C:\NDT\PJ\MediSign_AI - Copy")
BUILD = ROOT / "report_build"
FIG = BUILD / "figures"
OUT = ROOT / "BaoCao_NCKH_MediSignAI_NguyenDuyThuan_2311555799.docx"

# ---------- Constants -----------
FONT = "Times New Roman"
FONT_SIZE = Pt(12)
INDENT = Cm(1.27)
LINE_SPACING = 1.5
SPACE_BEFORE = Pt(6)
SPACE_AFTER = Pt(6)


# ============== Helpers ====================
def set_run_font(run, size=12, bold=False, italic=False, color=None, allcaps=False):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if allcaps:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "1")
        rPr.append(caps)


def set_paragraph_format(p, indent_first=True, align=None, space_before=None,
                         space_after=None, line=None, indent_left=None, keep_with_next=False):
    pf = p.paragraph_format
    pf.line_spacing = line if line is not None else LINE_SPACING
    pf.space_before = SPACE_BEFORE if space_before is None else space_before
    pf.space_after = SPACE_AFTER if space_after is None else space_after
    pf.first_line_indent = INDENT if indent_first else None
    if indent_left is not None:
        pf.left_indent = indent_left
    if align is not None:
        p.alignment = align
    if keep_with_next:
        pf.keep_with_next = True


def add_para(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent_first=True, indent_left=None, color=None, space_before=None,
             space_after=None, line=None, allcaps=False, keep_with_next=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color, allcaps=allcaps)
    set_paragraph_format(p, indent_first=indent_first, align=align,
                         space_before=space_before, space_after=space_after,
                         line=line, indent_left=indent_left, keep_with_next=keep_with_next)
    return p


def add_heading_1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    set_run_font(run, size=14, bold=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(18), space_after=Pt(12), keep_with_next=True)
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(12), space_after=Pt(6), keep_with_next=True)
    return p


def add_heading_3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, italic=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         indent_left=INDENT, space_before=Pt(8), space_after=Pt(4),
                         keep_with_next=True)
    return p


def add_heading_4(doc, text):
    p = doc.add_paragraph(style="Heading 4")
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         indent_left=INDENT, space_before=Pt(6), space_after=Pt(3),
                         keep_with_next=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=12)
    set_paragraph_format(p, indent_first=False, indent_left=INDENT,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    return p


def add_dash_item(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"-\t{text}")
    set_run_font(run, size=12)
    set_paragraph_format(p, indent_first=False, indent_left=INDENT,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    return p


def add_image(doc, path, width_cm=15.0, caption=None, source=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(6), space_after=Pt(2))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(caption)
        set_run_font(run, size=12, bold=True)
        set_paragraph_format(c, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(2), space_after=Pt(2))
    if source:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = s.add_run(source)
        set_run_font(run, size=11, italic=True)
        set_paragraph_format(s, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(0), space_after=Pt(8))


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, size=12, bold=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(6), space_after=Pt(2))


def add_table_source(doc, source):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(source)
    set_run_font(run, size=11, italic=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(8))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "D9E2F3")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shade)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=11)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    run._element.append(OxmlElement("w:lastRenderedPageBreak"))
    p2 = doc.add_paragraph()
    p2.add_run().add_break(6)  # WD_BREAK.PAGE = 6
    return p2


def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:br")
    fld.set(qn("w:type"), "page")
    r._element.append(fld)


def configure_section(section, page_num_format="decimal"):
    """Set margins, header/footer, page number format."""
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), page_num_format)
    pgNumType.set(qn("w:start"), "1")


def add_page_number(doc, section, restart=True):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = ""
    run = p.add_run()
    set_run_font(run, size=11)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(txt)
    run._element.append(fld_end)


def add_toc_field(doc, levels=3, title="MUC LUC"):
    p_title = add_heading_1(doc, title)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Vui long mo file trong Microsoft Word va bam phim F9 de cap nhat muc luc."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(placeholder)
    run._element.append(fld_end)
    set_run_font(run, size=12, italic=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(6), space_after=Pt(6))


def add_tof_field(doc, label="Hinh", title="DANH MUC HINH ANH"):
    add_heading_1(doc, title)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\h \\z \\c "{label}"'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mo file trong Word va F9 de cap nhat danh muc."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(placeholder)
    run._element.append(fld_end)
    set_run_font(run, size=12, italic=True)
    set_paragraph_format(p, indent_first=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(6), space_after=Pt(6))


def style_default(doc):
    """Apply Times New Roman 12 default to body."""
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = SPACE_BEFORE
    pf.space_after = SPACE_AFTER


print("helpers OK")
